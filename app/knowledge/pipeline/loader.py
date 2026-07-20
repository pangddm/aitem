from __future__ import annotations

from dataclasses import dataclass, field
from pathlib import Path
from typing import Any


@dataclass
class LoadedDocument:
    """DocumentLoader 的返回值"""

    text: str
    filename: str
    mime_type: str
    file_size: int
    metadata: dict[str, Any] = field(default_factory=dict)


class DocumentLoader:

    # MIME type mapping — 纯文本可直接 read_text
    TEXT_MIME_MAP = {
        ".txt": "text/plain",
        ".md": "text/markdown",
        ".log": "text/plain",
        ".json": "application/json",
        ".yaml": "text/yaml",
        ".yml": "text/yaml",
        ".csv": "text/csv",
        ".xml": "text/xml",
        ".html": "text/html",
        ".htm": "text/html",
        ".py": "text/x-python",
        ".sh": "text/x-shellscript",
        ".bat": "text/x-bat",
        ".ps1": "text/x-powershell",
        ".sql": "text/x-sql",
        ".cfg": "text/plain",
        ".conf": "text/plain",
        ".ini": "text/plain",
        ".toml": "text/toml",
    }

    SUPPORTED_SUFFIXES = sorted(
        set(TEXT_MIME_MAP) | {".docx", ".xlsx"}
    )

    async def load(
        self,
        file_path: str,
    ) -> LoadedDocument:

        path = Path(file_path)

        suffix = path.suffix.lower()

        if suffix not in self.SUPPORTED_SUFFIXES:
            raise ValueError(
                f"不支持的文件类型: {suffix}\n"
                f"支持的类型: {', '.join(self.SUPPORTED_SUFFIXES)}"
            )

        # ── docx ────────────────────────────────────────
        if suffix == ".docx":
            return self._load_docx(path)

        # ── xlsx ────────────────────────────────────────
        if suffix == ".xlsx":
            return self._load_xlsx(path)

        # ── 纯文本 ──────────────────────────────────────
        text = path.read_text(
            encoding="utf-8",
            errors="ignore",
        )

        mime_type = self.TEXT_MIME_MAP[suffix]
        file_size = path.stat().st_size

        return LoadedDocument(
            text=text,
            filename=path.name,
            mime_type=mime_type,
            file_size=file_size,
            metadata={
                "path": str(path.absolute()),
                "suffix": suffix,
            },
        )

    # ────────── docx 解析 ───────────────────────────────

    @staticmethod
    def _load_docx(path: Path) -> LoadedDocument:
        from docx import Document as DocxDocument

        doc = DocxDocument(str(path))

        parts = []
        for para in doc.paragraphs:
            text = para.text.strip()
            if text:
                style = para.style.name if para.style else ""
                if "Heading" in style or para.style and "Title" in style:
                    level = 1
                    if para.style and hasattr(para.style, 'name'):
                        for ch in para.style.name:
                            if ch.isdigit():
                                level = int(ch)
                                break
                    prefix = "#" * min(level, 6)
                    parts.append(f"\n{prefix} {text}")
                else:
                    parts.append(text)

        # 表格 → Markdown 表格
        for ti, table in enumerate(doc.tables):
            parts.append(f"\n### 表格 {ti + 1}\n")
            rows = []
            for row in table.rows:
                cells = [
                    cell.text.strip().replace("\n", " ")
                    for cell in row.cells
                ]
                rows.append(cells)

            if not rows:
                continue

            # 表头
            header = rows[0]
            parts.append("| " + " | ".join(header) + " |")
            parts.append("|" + "|".join("---" for _ in header) + "|")

            # 数据行（最多取前 50 行防过大）
            for row in rows[1:51]:
                parts.append("| " + " | ".join(row) + " |")

        text = "\n".join(parts)

        if not text.strip():
            text = path.read_text(encoding="utf-8", errors="ignore")

        return LoadedDocument(
            text=text,
            filename=path.name,
            mime_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
            file_size=path.stat().st_size,
            metadata={
                "path": str(path.absolute()),
                "suffix": ".docx",
            },
        )

    # ────────── xlsx 解析 ───────────────────────────────

    @staticmethod
    def _load_xlsx(path: Path) -> LoadedDocument:
        from openpyxl import load_workbook

        wb = load_workbook(
            str(path),
            read_only=True,
            data_only=True,
        )

        lines = []
        for sheet_name in wb.sheetnames:
            ws = wb[sheet_name]
            sheet_rows = [f"=== Sheet: {sheet_name} ==="]
            for row in ws.iter_rows():
                cells = [
                    str(cell.value) if cell.value is not None else ""
                    for cell in row
                ]
                line = "\t".join(cells).strip()
                if line:
                    sheet_rows.append(line)
            if len(sheet_rows) > 1:
                lines.extend(sheet_rows)
                lines.append("")

        text = "\n".join(lines)

        return LoadedDocument(
            text=text,
            filename=path.name,
            mime_type="application/vnd.openxmlformats-officedocument.spreadsheetml.sheet",
            file_size=path.stat().st_size,
            metadata={
                "path": str(path.absolute()),
                "suffix": ".xlsx",
            },
        )