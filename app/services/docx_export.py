"""将 Markdown 文本转换为 Word (.docx) 二进制内容。

供「下载报告」接口使用：把内部生成的 Markdown 报告转为排版良好的 Word 文档。
"""
from __future__ import annotations

import io
import re

from docx import Document
from docx.shared import Pt
from docx.oxml import OxmlElement
from docx.oxml.ns import qn


_INLINE_RE = re.compile(r"(\*\*[^*]+\*\*|`[^`]+`|\*[^*]+\*)")


def _set_cjk(run, bold=None, mono=False, size=None):
    """统一设置字体（含中文 East Asia 字体），避免中文在 Word 里变形。"""
    if size:
        run.font.size = Pt(size)
    if bold is not None:
        run.font.bold = bold
    run.font.name = "Consolas" if mono else "Microsoft YaHei"
    # 让中文使用中文字体（关键：仅设置 name 对中文无效）
    r_pr = run._element.get_or_add_rPr()
    r_fonts = r_pr.find(qn("w:rFonts"))
    if r_fonts is None:
        r_fonts = OxmlElement("w:rFonts")
        r_pr.append(r_fonts)
    r_fonts.set(qn("w:eastAsia"), "Microsoft YaHei")


def _add_inline(par, text):
    """把含 **加粗**、`行内代码`、*斜体* 的文本拆成多个 run。"""
    for token in _INLINE_RE.split(text):
        if not token:
            continue
        if token.startswith("**") and token.endswith("**") and len(token) > 2:
            run = par.add_run(token[2:-2])
            _set_cjk(run, bold=True)
        elif token.startswith("`") and token.endswith("`") and len(token) > 1:
            run = par.add_run(token[1:-1])
            _set_cjk(run, mono=True, size=10)
        elif token.startswith("*") and token.endswith("*") and not token.startswith("**"):
            run = par.add_run(token[1:-1])
            _set_cjk(run)
            run.font.italic = True
        else:
            _set_cjk(par.add_run(token))


def _split_row(line: str) -> list:
    s = line.strip()
    if s.startswith("|"):
        s = s[1:]
    if s.endswith("|"):
        s = s[:-1]
    return [c.strip() for c in s.split("|")]


def _add_hr(doc) -> None:
    p = doc.add_paragraph()
    p_pr = p._p.get_or_add_pPr()
    p_bdr = OxmlElement("w:pBdr")
    bottom = OxmlElement("w:bottom")
    bottom.set(qn("w:val"), "single")
    bottom.set(qn("w:sz"), "6")
    bottom.set(qn("w:space"), "1")
    bottom.set(qn("w:color"), "999999")
    p_bdr.append(bottom)
    p_pr.append(p_bdr)


def markdown_to_docx(md_text: str) -> io.BytesIO:
    """把 markdown 文本转成 .docx，返回 BytesIO。"""
    doc = Document()

    # 默认正文样式：中文字体
    normal = doc.styles["Normal"]
    normal.font.name = "Microsoft YaHei"
    normal.font.size = Pt(11)
    r_fonts = normal.element.get_or_add_rPr().get_or_add_rFonts()
    r_fonts.set(qn("w:eastAsia"), "Microsoft YaHei")

    lines = md_text.split("\n")
    n = len(lines)
    i = 0
    while i < n:
        raw = lines[i]
        s = raw.rstrip()
        stripped = s.strip()

        # 我们自己生成的 <details>/<summary> 标签
        if stripped in ("<details>", "</details>"):
            i += 1
            continue
        if stripped.startswith("<summary>"):
            title = stripped.replace("<summary>", "").replace("</summary>", "").strip()
            p = doc.add_heading("", level=4)
            _add_inline(p, title)
            i += 1
            continue

        # 分隔线
        if re.match(r"^\s*(-{3,}|\*{3,})\s*$", s):
            _add_hr(doc)
            i += 1
            continue

        # 代码块
        if re.match(r"^\s*```", s):
            buf = []
            i += 1
            while i < n and not re.match(r"^\s*```", lines[i]):
                buf.append(lines[i])
                i += 1
            i += 1  # 跳过收尾 ```
            for cl in buf:
                _set_cjk(doc.add_paragraph().add_run(cl), mono=True, size=9.5)
            continue

        # 标题
        m = re.match(r"^(#{1,6})\s+(.*)$", s)
        if m:
            level = len(m.group(1))
            p = doc.add_heading("", level=min(level, 6))
            _add_inline(p, m.group(2))
            i += 1
            continue

        # 表格（| a | b |  分隔行）
        if stripped.startswith("|") and i + 1 < n and re.match(r"^\s*\|[\s:|-]+\|\s*$", lines[i + 1]):
            header = _split_row(s)
            rows = []
            i += 2
            while i < n and lines[i].strip().startswith("|"):
                rows.append(_split_row(lines[i]))
                i += 1
            ncol = len(header)
            table = doc.add_table(rows=1 + len(rows), cols=ncol)
            table.style = "Light Grid Accent 1"
            for c, htext in enumerate(header):
                _add_inline(table.rows[0].cells[c].paragraphs[0], htext)
            for r_i, row in enumerate(rows, start=1):
                for c in range(min(ncol, len(row))):
                    _add_inline(table.rows[r_i].cells[c].paragraphs[0], row[c])
            continue

        # 引用
        if stripped.startswith(">"):
            body = s.lstrip().lstrip(">").strip()
            p = doc.add_paragraph()
            p.paragraph_format.left_indent = Pt(18)
            _add_inline(p, body)
            i += 1
            continue

        # 无序列表
        m = re.match(r"^\s*[-*]\s+(.*)$", s)
        if m:
            p = doc.add_paragraph(style="List Bullet")
            _add_inline(p, m.group(1))
            i += 1
            continue

        # 空行
        if stripped == "":
            i += 1
            continue

        # 普通段落
        p = doc.add_paragraph()
        _add_inline(p, s)
        i += 1

    buf = io.BytesIO()
    doc.save(buf)
    buf.seek(0)
    return buf